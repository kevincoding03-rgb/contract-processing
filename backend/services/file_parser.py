import io
from PyPDF2 import PdfReader
from docx import Document
import fitz  # PyMuPDF
from PIL import Image


# 太小的图片（如图标、签名线条）跳过，避免无意义的 OCR 调用
_MIN_IMAGE_AREA = 64 * 64  # 像素


def _to_png_bytes(img_bytes: bytes) -> bytes:
    """把任意格式的图片字节统一转成 PNG，方便后续 base64 编码和视觉模型处理。"""
    try:
        img = Image.open(io.BytesIO(img_bytes))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        # 转换失败就返回原始字节，交给视觉模型自己处理
        return img_bytes


def _is_meaningful_image(img_bytes: bytes) -> bool:
    """过滤掉过小的图标/装饰图。"""
    try:
        img = Image.open(io.BytesIO(img_bytes))
        w, h = img.size
        return w * h >= _MIN_IMAGE_AREA
    except Exception:
        return False


def parse_txt(content: bytes) -> dict:
    return {"text": content.decode("utf-8", errors="ignore"), "images": []}


def parse_pdf(content: bytes) -> dict:
    """用 PyMuPDF 提取文字 + 内嵌图片。

    PyMuPDF 对扫描版 PDF、混合排版 PDF 的支持远好于 PyPDF2。
    """
    text_parts = []
    images = []

    doc = fitz.open(stream=content, filetype="pdf")
    for page_index in range(len(doc)):
        page = doc[page_index]

        # 1) 提取文字
        page_text = page.get_text("text")
        if page_text and page_text.strip():
            text_parts.append(page_text.strip())

        # 2) 提取内嵌图片
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image.get("image")
                if img_bytes and _is_meaningful_image(img_bytes):
                    images.append({
                        "source": f"PDF 第{page_index + 1}页",
                        "data": _to_png_bytes(img_bytes),
                    })
            except Exception:
                continue

    doc.close()
    return {"text": "\n".join(text_parts), "images": images}


def parse_docx(content: bytes) -> dict:
    """提取 Word 段落文字 + 内嵌图片。"""
    doc = Document(io.BytesIO(content))

    # 1) 段落文字
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

    # 2) 表格里的文字（合同常用表格）
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    # 3) 内嵌图片（来自 inline shapes，即正文插入的图片）
    images = []
    inline_shapes = getattr(doc, "inline_shapes", None)
    if inline_shapes is not None:
        for idx, shape in enumerate(inline_shapes):
            try:
                # python-docx 没有直接给 bytes 的 API，需要从底层关系取
                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                rId = blip.embed
                image_part = doc.part.related_parts[rId]
                img_bytes = image_part.blob
                if img_bytes and _is_meaningful_image(img_bytes):
                    images.append({
                        "source": f"Word 第{idx + 1}张图片",
                        "data": _to_png_bytes(img_bytes),
                    })
            except Exception:
                continue

    return {"text": "\n".join(text_parts), "images": images}


def parse_file(filename: str, content: bytes) -> dict:
    """解析文件，返回 {'text': str, 'images': [{'source': str, 'data': bytes}, ...]}。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parsers = {
        "txt": parse_txt,
        "pdf": parse_pdf,
        "docx": parse_docx,
    }
    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件格式: .{ext}，仅支持 txt、pdf、docx")
    return parser(content)
